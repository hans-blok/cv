#!/usr/bin/env python3
"""
Generate DOCX from CV data with similar layout to PDF.

Reads directly from docs/data/**/*.yml (plain content, no HTML/CSS
round-tripping) - the same source of truth the MkDocs site itself is
built from via main.py's macros.
"""

from pathlib import Path
import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Paths
ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "docs" / "data"
ENGAGEMENTS_DIR = DATA_DIR / "engagements"
OUTPUT_FILE = ROOT / "docs" / "assets" / "cv.docx"

# Colors (matching docs/stylesheets/style.css)
TEXT_PRIMARY = RGBColor(0x11, 0x11, 0x11)  # #111 - var(--text-primary)
GREY_TEXT = RGBColor(99, 110, 114)   # #636e72 - var(--text-muted)
GREY_LIGHT = RGBColor(178, 190, 195)  # #b2bec3 - var(--sep-block)
TAG_BG = "DFE6E9"  # #dfe6e9 - var(--sep-course), same background as the site's expertise-tag chips

CONTACT_ICONS = {
    "linkedin": "🔗",
    "website": "🌐",
    "github": "💻",
    "phone": "☎",
    "email": "✉",
}

def load_yaml(name):
    path = DATA_DIR / name
    if not path.is_file():
        return None
    return yaml.safe_load(path.read_text(encoding='utf-8'))

def load_engagements():
    """All engagement records, newest first - same sort key as main.py's macro."""
    if not ENGAGEMENTS_DIR.is_dir():
        return []
    files = sorted(
        ENGAGEMENTS_DIR.glob("*.yml"),
        key=lambda p: yaml.safe_load(p.read_text(encoding='utf-8'))["order"],
        reverse=True,
    )
    return [yaml.safe_load(f.read_text(encoding='utf-8')) for f in files]

def collect_expertise_tags(engagements):
    """Deduplicated tag list from engagement keywords + certification names -
    same algorithm as main.py's render_expertise_tags(), so the DOCX shows the
    same Kernexpertise content as the site/PDF."""
    seen = set()
    tags = []

    def add(raw):
        tag = (raw or "").strip()
        key = tag.lower()
        if tag and key not in seen:
            seen.add(key)
            tags.append(tag)

    for engagement in engagements:
        for kw in (engagement.get("keywords") or "").split(","):
            add(kw)
    for cert in load_yaml("certifications.yml") or []:
        add(cert.get("name", ""))
    return tags

def md_bullets_to_docx_text(text):
    """Normalize migrated content ('- ' Markdown bullets, trailing hard-break
    spaces) into the plain '• '-bulleted, line-broken text add_text_with_breaks()
    expects - the same bullet character the site/PDF renders via the |markdown
    Jinja filter turning '- ' into <li>, here turned into '• ' directly."""
    if not text:
        return ""
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("- "):
            line = "• " + line[2:]
        lines.append(line)
    return "\n".join(lines)

def add_horizontal_line(doc):
    """Add a horizontal line separator"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B2BEC3')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_shading(cell, hex_color):
    """Fill a table cell with a solid background color"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_title(doc, title):
    """Add section title (dark, uppercase, bold)"""
    p = doc.add_paragraph()
    p.add_run(title.upper()).font.color.rgb = TEXT_PRIMARY
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(11)
    p.space_after = Pt(6)
    return p

def add_hyperlink(paragraph, url, text, color="1155CC", size_pt=None):
    """Add a clickable hyperlink run to a paragraph (python-docx has no
    built-in API for this - standard low-level OOXML recipe)."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))  # half-points
        rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_tag_grid(doc, tags, columns=4):
    """Render tags as a borderless, shaded-cell grid - approximates the site's
    flex-wrap 'chip' look far better than one long comma-separated paragraph."""
    if not tags:
        return
    rows = (len(tags) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.autofit = True
    for idx, tag in enumerate(tags):
        r, c = divmod(idx, columns)
        cell = table.rows[r].cells[c]
        set_cell_border(cell, top=True, left=True, bottom=True, right=True)
        set_cell_shading(cell, TAG_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(tag)
        run.font.size = Pt(9.5)
    # Fill any leftover cells in the last row with empty shaded space (keeps a tidy grid)
    for idx in range(len(tags), rows * columns):
        r, c = divmod(idx, columns)
        set_cell_shading(table.rows[r].cells[c], "FFFFFF")

def add_detail_label(doc, text):
    """Small grey/bold sub-label, e.g. 'Werkzaamheden' or 'Opleidingen'."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = GREY_TEXT
    run.font.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_text_with_breaks(paragraph, text):
    """Add text with line breaks preserved"""
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if line.strip():
            paragraph.add_run(line.strip())
            if i < len(lines) - 1:
                paragraph.add_run('\n')

def generate_docx():
    """Generate Word document"""
    # Check whether the output file is locked by another process (e.g. Word)
    if OUTPUT_FILE.exists():
        try:
            with OUTPUT_FILE.open('r+b'):
                pass
        except PermissionError:
            print(f"Error: {OUTPUT_FILE} is locked by another application (e.g. Word).")
            print("Sluit het Word-document en probeer opnieuw.")
            return False

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Personal Data
    personal = load_yaml("personal-data.yml") or {}
    fields = personal.get("fields", [])

    add_section_title(doc, 'PERSOONLIJKE GEGEVENS')

    # Photo, right-aligned (mirrors the site's profile-photo placement)
    # personal-data.yml's `photo` path is relative to docs_dir (e.g. "assets/img/..."),
    # same as how it's referenced in the built HTML.
    photo_path = DATA_DIR.parent / personal.get("photo", "")
    if personal.get("photo") and photo_path.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(photo_path), width=Inches(1.2))
        p.paragraph_format.space_after = Pt(6)

    # Compact contact line: icon + short label only (never the raw URL),
    # clickable via a real hyperlink - site's sidebar has no direct DOCX
    # equivalent, but a recruiter reading a downloaded Word file still needs
    # a way to reach out.
    contacts = load_yaml("contact.yml") or []
    if contacts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        for i, item in enumerate(contacts):
            if i > 0:
                sep = p.add_run("     ")
                sep.font.size = Pt(9.5)
            icon = p.add_run(CONTACT_ICONS.get(item.get("type"), "") + " ")
            icon.font.size = Pt(9.5)
            add_hyperlink(p, item.get("url", ""), item.get("label", ""), size_pt=9.5)

    if fields:
        table = doc.add_table(rows=len(fields), cols=2)
        table.autofit = False
        table.allow_autofit = False

        for idx, f in enumerate(fields):
            row = table.rows[idx]
            label_cell = row.cells[0]
            value_cell = row.cells[1]

            set_cell_border(label_cell, top=True, left=True, bottom=True, right=True)
            set_cell_border(value_cell, top=True, left=True, bottom=True, right=True)

            p = label_cell.paragraphs[0]
            run = p.add_run(f.get("label", ""))
            run.font.color.rgb = GREY_TEXT
            run.font.bold = True
            run.font.size = Pt(11)

            value_cell.text = f.get("value", "")
            value_cell.paragraphs[0].runs[0].font.size = Pt(11)

    add_horizontal_line(doc)

    # Kernexpertise (deduplicated tags - same source/order as the site's Kernexpertise block)
    engagements = load_engagements()
    tags = collect_expertise_tags(engagements)

    if tags:
        add_section_title(doc, 'KERNEXPERTISE')
        add_tag_grid(doc, tags)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        add_horizontal_line(doc)

    # Personal Text
    personal_text_file = DATA_DIR / "personal-text.md"
    text = personal_text_file.read_text(encoding='utf-8') if personal_text_file.is_file() else ""
    if text:
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                add_text_with_breaks(p, para.strip())
                p.paragraph_format.space_after = Pt(6)

    add_horizontal_line(doc)

    # Achtergrond (Opleidingen + Belangrijkste certificeringen, merged like the site's
    # collapsed "Achtergrond" block - always shown here since DOCX has no accordion)
    add_section_title(doc, 'ACHTERGROND')

    educations = load_yaml("educations.yml") or []
    if educations:
        add_detail_label(doc, "Opleidingen")
        for row in educations:
            p = doc.add_paragraph()
            run = p.add_run(f"{row.get('period', '')}  ")
            run.font.color.rgb = GREY_TEXT
            run.font.size = Pt(11)
            run = p.add_run(row.get("name", ""))
            if row.get("institute"):
                run = p.add_run(f" — {row['institute']}")
            if row.get("place"):
                run = p.add_run(f", {row['place']}")
            p.paragraph_format.space_after = Pt(3)

    certifications = load_yaml("certifications.yml") or []
    if certifications:
        add_detail_label(doc, "Belangrijkste certificeringen")
        for row in certifications:
            p = doc.add_paragraph()
            run = p.add_run(f"{row.get('period', '')}  ")
            run.font.color.rgb = GREY_TEXT
            run = p.add_run(row.get("name", ""))
            if row.get("institute"):
                run = p.add_run(f" — {row['institute']}")
            p.paragraph_format.space_after = Pt(3)

    add_horizontal_line(doc)

    # Courses
    add_section_title(doc, 'CURSUSSEN')
    courses = load_yaml("courses.yml") or []

    for group in courses:
        p = doc.add_paragraph()
        run = p.add_run(f"{group.get('period', '')}  ")
        run.font.color.rgb = GREY_TEXT
        run = p.add_run(group.get("items_text", ""))
        p.paragraph_format.space_after = Pt(3)

    add_horizontal_line(doc)

    # Overige cursussen
    courses_short = load_yaml("courses-short.yml") or []
    if courses_short:
        add_section_title(doc, 'OVERIGE CURSUSSEN')
        for group in courses_short:
            p = doc.add_paragraph()
            if group.get('period'):
                run = p.add_run(f"{group['period']}  ")
                run.font.color.rgb = GREY_TEXT
            run = p.add_run(group.get("items_text", ""))
            p.paragraph_format.space_after = Pt(3)
        add_horizontal_line(doc)

    # Werkervaring (Engagements)
    add_section_title(doc, 'WERKERVARING')

    if engagements:
        for eng in engagements:
            # Summary line
            p = doc.add_paragraph()
            if eng.get("period"):
                run = p.add_run(f"{eng['period']}  ")
                run.font.color.rgb = GREY_TEXT
            if eng.get("organisation"):
                run = p.add_run(eng["organisation"])
            if eng.get("role"):
                run = p.add_run(f" — {eng['role']}")
            p.paragraph_format.space_after = Pt(6)

            # Details
            if eng.get("activities"):
                add_detail_label(doc, "Werkzaamheden")
                p = doc.add_paragraph()
                add_text_with_breaks(p, md_bullets_to_docx_text(eng["activities"]))
                p.paragraph_format.space_after = Pt(6)

            if eng.get("achievements"):
                add_detail_label(doc, "Belangrijkste prestaties")
                p = doc.add_paragraph()
                add_text_with_breaks(p, md_bullets_to_docx_text(eng["achievements"]))
                p.paragraph_format.space_after = Pt(6)

            if eng.get("keywords"):
                add_detail_label(doc, "Trefwoorden")
                p = doc.add_paragraph()
                add_text_with_breaks(p, eng["keywords"])
                p.paragraph_format.space_after = Pt(6)

            # Separator between engagements
            add_horizontal_line(doc)

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"✓ Word document generated: {OUTPUT_FILE}")

if __name__ == "__main__":
    try:
        generate_docx()
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
